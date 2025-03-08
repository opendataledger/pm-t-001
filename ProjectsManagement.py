import streamlit as st
import pandas as pd
from datetime import date, datetime
import os
import altair as alt
import smtplib
from email.mime.text import MIMEText
import base64

# import streamlit as st
# import pandas as pd
from docx import Document
# import smtplib
# from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
# from email.mime.base import MIMEBase
from email import encoders
import os

# import smtplib
# from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart

# File to store project data
PROJECT_DATA_FILE = 'projects.csv'
TASK_DATA_FILE = 'tasks.csv'
TEAM_DATA_FILE = 'team_members.csv'

st.title("Projects Management Tool")
# st.title("Developed By Mansoor Sarookh, CS Student at GPGC Swabi")

# Load project data
def load_projects():
    try:
        return pd.read_csv(PROJECT_DATA_FILE)
    except FileNotFoundError:
        return pd.DataFrame(columns=["Project Name", "Description", "Start Date", "End Date", "Priority", "Status", "Progress", "Budget", "Tags"])

# Save project data
def save_projects(df):
    df.to_csv(PROJECT_DATA_FILE, index=False)

# Load tasks
def load_tasks():
    try:
        return pd.read_csv(TASK_DATA_FILE)
    except FileNotFoundError:
        return pd.DataFrame(columns=["Project Name", "Task Name", "Deadline", "Status", "Assigned To"])

# Save tasks
def save_tasks(df):
    df.to_csv(TASK_DATA_FILE, index=False)

# Load team members
def load_team():
    try:
        return pd.read_csv(TEAM_DATA_FILE)
    except FileNotFoundError:
        return pd.DataFrame(columns=["Project Name", "Member Name", "Role"])

# Save team members
def save_team(df):
    df.to_csv(TEAM_DATA_FILE, index=False)






import streamlit as st

def home():
    st.title("Welcome to the Project Management Tool")
    st.markdown("""
    ## About the Tool
    The **Project Management Tool** is a comprehensive solution built specifically for **Admins** and **Project Managers** in small software houses.
                - It **streamlines workflows, enhances productivity, and ensures smooth project execution.**

    ### Key Features:
    - **Add, Edit, and Delete Projects:** Easily create, modify, or remove projects.
    - **Task Management:** Assign and organize tasks efficiently.
    - **High Priority & Deadlines:** Focus on critical tasks to meet deadlines.
    - **Project Progress Tracking:** Monitor project progress in real-time.
    - **Team Collaboration:** Facilitate communication and teamwork.
    - **Analytics & Reports:** Generate performance and resource utilization reports.
    - **Calendar View:** Plan schedules with an intuitive calendar format.
    - **Search and Filter Projects:** Quickly find projects using advanced search tools.
    - **Budget Tracking:** Manage budgets effectively to avoid overspending.
    - **Comments and Notes:** Document critical details and discussions.
    - **Prioritization Metrics:** Focus on the most impactful tasks.
    - **Gantt Chart Visualization:** Gain a clear view of timelines with Gantt charts.
    - **Resource Management:** Allocate resources optimally across projects.
    - **Risk Management:** Identify and mitigate potential risks proactively.

    ### Why Choose This Tool?
    - Specifically designed for **Admins** and **Project Managers**.
    - Tailored for small software houses.
    - Focused on streamlining project workflows and improving efficiency.
    """)



# Add new project with suggestions and technology selection
def add_project():
    st.header("Add New Project")

    # Project name input with suggestions
    project_suggestions = [
        "Machine Learning Model",
        "Web Development App",
        "Mobile App Development",
        "Data Science Pipeline",
        "API Development",
        "Game Development",
    ]
    project_name = st.selectbox("Select or type a project name:", project_suggestions)

    # Description based on project name
    description_placeholder = {
        "Machine Learning Model": "Enter a brief description of your machine learning model project.",
        "Web Development App": "Enter a brief description of your web development app project.",
        "Mobile App Development": "Enter a brief description of your mobile app development project.",
        "Data Science Pipeline": "Enter a brief description of your data science pipeline project.",
        "API Development": "Enter a brief description of your API development project.",
        "Game Development": "Enter a brief description of your game development project."
    }
    description = st.text_area("Project Description", placeholder=description_placeholder.get(project_name, ""))

    # Technology selection dynamically based on project name
    technologies = {
        "Machine Learning Model": ["Python", "TensorFlow", "PyTorch", "Scikit-learn"],
        "Web Development App": ["HTML", "CSS", "JavaScript", "Django", "Flask"],
        "Mobile App Development": ["Flutter", "React Native", "Swift", "Kotlin"],
        "Data Science Pipeline": ["Python", "Pandas", "NumPy", "Matplotlib"],
        "API Development": ["FastAPI", "Flask", "Express.js", "Django Rest Framework"],
        "Game Development": ["Unity", "Unreal Engine", "C#", "Godot"],
    }

    if project_name in technologies:
        tech_options = technologies[project_name]
    else:
        tech_options = []
    
    technology_used = st.multiselect("Select technologies used in the project:", tech_options)

    # Additional project details
    start_date = st.date_input("Start Date", date.today())
    end_date = st.date_input("End Date")
    priority = st.selectbox("Priority", ["Low", "Medium", "High"])
    budget = st.number_input("Budget", min_value=0.0)

    # Submit button
    if st.button("Create Project"):
        if project_name and description:
            new_project = pd.DataFrame({
                "Project Name": [project_name],
                "Description": [description],
                "Start Date": [start_date],
                "End Date": [end_date],
                "Priority": [priority],
                "Status": ["Uncompleted"],
                "Progress": [0],
                "Budget": [budget],
                "Technologies": [", ".join(technology_used)]  # Store selected technologies as a comma-separated string
            })
            df = load_projects()
            df = pd.concat([df, new_project], ignore_index=True)
            save_projects(df)
            st.success(f"Project '{project_name}' has been added!")
        else:
            st.error("Please fill in all fields")

# Edit project with suggestions and technology update
def edit_project():
    st.header("Edit Project")
    df = load_projects()
    if df.empty:
        st.info("No projects to edit.")
        return
    
    project_names = df["Project Name"].tolist()
    project = st.selectbox("Select Project", project_names)
    
    if project:
        row = df[df["Project Name"] == project].iloc[0]

        # Project name input with suggestions
        project_suggestions = [
            "Machine Learning Model",
            "Web Development App",
            "Mobile App Development",
            "Data Science Pipeline",
            "API Development",
            "Game Development",
        ]
        new_name = st.selectbox("Select or type a project name:", project_suggestions, index=project_suggestions.index(row["Project Name"]) if row["Project Name"] in project_suggestions else 0)

        # Description based on project name
        description_placeholder = {
            "Machine Learning Model": "Enter a brief description of your machine learning model project.",
            "Web Development App": "Enter a brief description of your web development app project.",
            "Mobile App Development": "Enter a brief description of your mobile app development project.",
            "Data Science Pipeline": "Enter a brief description of your data science pipeline project.",
            "API Development": "Enter a brief description of your API development project.",
            "Game Development": "Enter a brief description of your game development project."
        }
        description = st.text_area("Project Description", placeholder=description_placeholder.get(new_name, ""), value=row["Description"])

        # Technology selection dynamically based on project name
        technologies = {
            "Machine Learning Model": ["Python", "TensorFlow", "PyTorch", "Scikit-learn"],
            "Web Development App": ["HTML", "CSS", "JavaScript", "Django", "Flask"],
            "Mobile App Development": ["Flutter", "React Native", "Swift", "Kotlin"],
            "Data Science Pipeline": ["Python", "Pandas", "NumPy", "Matplotlib"],
            "API Development": ["FastAPI", "Flask", "Express.js", "Django Rest Framework"],
            "Game Development": ["Unity", "Unreal Engine", "C#", "Godot"],
        }

        tech_options = technologies.get(new_name, [])
        technology_used = st.multiselect("Select technologies used in the project:", tech_options, default=row["Technologies"].split(", "))

        # Additional project details
        start_date = st.date_input("Start Date", row["Start Date"])
        end_date = st.date_input("End Date", row["End Date"])
        priority = st.selectbox("Priority", ["Low", "Medium", "High"], index=["Low", "Medium", "High"].index(row["Priority"]))
        budget = st.number_input("Budget", min_value=0.0, value=row["Budget"])

        # Submit button
        if st.button("Update Project"):
            # Update the project data
            df.loc[df["Project Name"] == project, ["Project Name", "Description", "Start Date", "End Date", "Priority", "Budget", "Technologies"]] = new_name, description, start_date, end_date, priority, budget, ", ".join(technology_used)
            save_projects(df)
            st.success(f"Project '{new_name}' has been updated.")

# Delete project
def delete_project():
    st.header("Delete Project")
    df = load_projects()
    if df.empty:
        st.info("No projects to delete.")
        return
    
    project = st.selectbox("Select Project to Delete", df["Project Name"])
    if st.button("Delete Project"):
        df = df[df["Project Name"] != project]
        save_projects(df)
        st.success(f"Project '{project}' has been deleted.")

# THIS IS THE EXACT DESIRE ONE EXCEPT OF THE EMAIL FIELD ALTHOGUH WE CAN SEDN EMAIL SMOOTHLY

# Function to send email
def send_email(to_email, subject, body, attachment=None):
    from_email = "your_email@gmail.com"  # Replace with your email
    password = "your_password"          # Replace with your email app password
    try:
        server = smtplib.SMTP('smtp.gmail.com', 587)
        server.starttls()
        server.login(from_email, password)
        msg = MIMEMultipart()
        msg['Subject'] = subject
        msg['From'] = from_email
        msg['To'] = to_email
        msg.attach(MIMEText(body, 'plain'))

        # Attach a file if provided
        if attachment:
            with open(attachment, "rb") as file:
                part = MIMEBase("application", "octet-stream")
                part.set_payload(file.read())
            encoders.encode_base64(part)
            part.add_header("Content-Disposition", f"attachment; filename={attachment}")
            msg.attach(part)

        server.sendmail(from_email, to_email, msg.as_string())
        server.quit()
    except Exception as e:
        print(f"Failed to send email: {e}")

# Task management within each project
def manage_tasks():
    st.header("Task Management")

    # Load projects and tasks
    projects = load_projects()
    tasks_df = load_tasks()

    if projects.empty:
        st.info("No projects available.")
        return

    # Select a project
    project = st.selectbox("Select Project", projects["Project Name"])
    
    # Define team members
    team_members = {
        "Machine Learning Model": ["Ahmed", "Omar", "Faisal"],
        "Web Development App": ["Zaid", "Salman", "Yusuf"],
        "Mobile App Development": ["Hassan", "Bilal", "Ibrahim"],
        "Data Science Pipeline": ["Suleiman", "Imran", "Farhan"],
        "API Development": ["Adeel", "Hamza", "Mustafa"],
        "Game Development": ["Ali", "Khalid", "Tariq"]
    }

    # Task details
    st.subheader("Add a New Task")
    title = st.text_input("Task Title")
    description = st.text_area("Task Description")
    priority = st.selectbox("Priority", ["Low", "Medium", "High", "Custom"])
    custom_priority = None
    if priority == "Custom":
        custom_priority = st.text_input("Enter Custom Priority")

    tags = st.multiselect("Tags/Labels", ["Design", "Development", "Testing", "Research", "Bug Fixing"])
    team_leader = st.selectbox("Team Leader", team_members.get(project, []))
    assigned_to = st.multiselect("Assigned Team Members", team_members.get(project, []))
    start_date = st.date_input("Start Date")  # Start Date feature
    deadline = st.date_input("Deadline")

    # Prepare new task as a DataFrame (temporary)
    new_task = pd.DataFrame({
        "Project Name": [project],
        "Task Title": [title],
        "Description": [description],
        "Priority": [priority if priority != "Custom" else custom_priority],
        "Tags": [", ".join(tags)],
        "Team Leader": [team_leader],
        "Assigned To": [", ".join(assigned_to)],
        "Start Date": [start_date],  # Include Start Date
        "Deadline": [deadline],
        "Status": ["Uncompleted"]
    })

    # Show options before adding the task
    if title and team_leader and assigned_to:
        st.subheader("Preview Current Task")
        st.write(new_task)

        # Download Current Task
        st.subheader("Download Current Task")
        current_task_file = "current_task.csv"
        new_task.to_csv(current_task_file, index=False)
        st.download_button(
            label="Download Current Task as CSV",
            data=open(current_task_file, "rb").read(),
            file_name="current_task.csv",
            mime="text/csv"
        )

        # Send Current Task File via Email
        st.subheader("Send Email to Team Member")
        email_recipient = st.text_input("Email Recipient for Current Task File")
        if st.button("Send Current Task File"):
            if email_recipient:
                send_email(
                    to_email=email_recipient,
                    subject=f"Task Assigned: {title}",
                    body=f"Task '{title}' has been assigned to you. Please find the details attached.",
                    attachment=current_task_file
                )
                st.success(f"Current task file sent to {email_recipient}!")
            else:
                st.error("Please provide an email recipient.")

    # Add task button
    if st.button("Add Task"):
        if project and title and team_leader and assigned_to and start_date:
            tasks_df = pd.concat([tasks_df, new_task], ignore_index=True)
            save_tasks(tasks_df)
            st.success(f"Task '{title}' has been added to project '{project}'.")
        else:
            st.error("Please fill in all required fields.")




# View project deadlines and high-priority projects
def view_high_priority():
    st.header("Upcoming Deadlines & High Priority Projects")
    df = load_projects()
    if df.empty:
        st.info("No projects available.")
        return
    
    today = date.today()
    upcoming_deadlines = df[df["End Date"] >= str(today)].sort_values("End Date")
    high_priority = upcoming_deadlines[upcoming_deadlines["Priority"] == "High"]

    if not high_priority.empty:
        st.subheader("High Priority Projects")
        for i, row in high_priority.iterrows():
            st.text(f"{row['Project Name']} - Deadline: {row['End Date']}")

    st.subheader("Upcoming Deadlines")
    for i, row in upcoming_deadlines.iterrows():
        st.text(f"{row['Project Name']} - Deadline: {row['End Date']}")

# Team collaboration and management
def manage_team():
    st.header("Team Management")
    df = load_team()
    project = st.selectbox("Assign Team Member to Project", load_projects()["Project Name"])
    member_name = st.text_input("Team Member Name")
    role = st.selectbox("Role", ["Project Manager", "Developer", "Designer", "Quality Assurance Specialist QA"])

    if st.button("Add Member"):
        if project and member_name:
            new_member = pd.DataFrame({
                "Project Name": [project],
                "Member Name": [member_name],
                "Role": [role]
            })
            df = pd.concat([df, new_member], ignore_index=True)
            save_team(df)
            st.success(f"Team member '{member_name}' assigned to project '{project}'.")

    if not df.empty:
        st.subheader("Team Members")
        for i, member in df.iterrows():
            st.text(f"Member: {member['Member Name']} | Role: {member['Role']} | Project: {member['Project Name']}")





# Predefined Data
project_suggestions = [
    "Machine Learning Model",
    "Web Development App",
    "Mobile App Development",
    "Data Science Pipeline",
    "API Development",
    "Game Development",
]

technologies = {
    "Machine Learning Model": ["Python", "TensorFlow", "PyTorch", "Scikit-learn"],
    "Web Development App": ["HTML", "CSS", "JavaScript", "Django", "Flask"],
    "Mobile App Development": ["Flutter", "React Native", "Swift", "Kotlin"],
    "Data Science Pipeline": ["Python", "Pandas", "NumPy", "Matplotlib"],
    "API Development": ["FastAPI", "Flask", "Express.js", "Django Rest Framework"],
    "Game Development": ["Unity", "Unreal Engine", "C#", "Godot"],
}

# Function to generate project report
def generate_report(project_name, deadline, progress_data, overall_progress):
    doc = Document()
    doc.add_heading("Project Report", level=1)
    
    # Project Details
    doc.add_heading("Project Details", level=2)
    doc.add_paragraph(f"Project Name: {project_name}")
    doc.add_paragraph(f"Deadline: {deadline}")
    doc.add_paragraph(f"Overall Progress: {overall_progress:.2f}%")
    
    # Team Progress
    doc.add_heading("Team Progress", level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Team Member"
    hdr_cells[1].text = "Progress (%)"
    for progress in progress_data:
        row_cells = table.add_row().cells
        row_cells[0].text = progress["Member"]
        row_cells[1].text = f"{progress['Progress (%)']}%"
    
    # Save Report
    file_name = f"{project_name.replace(' ', '_')}_Report.docx"
    doc.save(file_name)
    return file_name

# Function to allow downloading the file
def create_download_link(file_path):
    with open(file_path, "rb") as file:
        file_bytes = file.read()
    b64 = base64.b64encode(file_bytes).decode()
    href = f'<a href="data:application/octet-stream;base64,{b64}" download="{os.path.basename(file_path)}">Download Report</a>'
    return href

# Function to send email with the report
def send_email_with_report(receiver_email, subject, body, file_path):
    sender_email = "your_email@example.com"  # Replace with your email
    msg = MIMEMultipart()
    msg["From"] = sender_email
    msg["To"] = receiver_email
    msg["Subject"] = subject

    # Email body
    msg.attach(MIMEText(body, "plain"))

    # Attach the document
    with open(file_path, "rb") as attachment:
        part = MIMEBase("application", "octet-stream")
        part.set_payload(attachment.read())
        encoders.encode_base64(part)
        part.add_header(
            "Content-Disposition",
            f"attachment; filename= {os.path.basename(file_path)}",
        )
        msg.attach(part)
    
    # Send email
    try:
        with smtplib.SMTP("smtp.gmail.com", 587) as server:
            server.starttls()
            server.sendmail(sender_email, receiver_email, msg.as_string())
        return f"Email sent successfully to {receiver_email}!"
    except Exception as e:
        return f"Failed to send email: {e}"

# Function to track project progress
def project_progress():
    st.title("Project Progress Tracking")
    
    # Select a project
    project = st.selectbox("Select a Project", project_suggestions)
    technologies_used = ", ".join(technologies[project])
    st.write(f"Technologies: {technologies_used}")
    
    # Input for deadline and progress
    deadline = st.date_input(f"Set a deadline for {project}")
    progress_data = []
    team_size = st.number_input(f"Enter number of team members for {project}", min_value=1, step=1)
    for i in range(team_size):
        member = st.text_input(f"Enter team member {i+1} name")
        progress = st.slider(f"Progress by {member} (%)", 0, 100, 0)
        email = st.text_input(f"Email for {member}")
        progress_data.append({"Member": member, "Progress (%)": progress, "Email": email})
    
    # Calculate overall progress
    progress_df = pd.DataFrame(progress_data)
    overall_progress = progress_df["Progress (%)"].mean()
    st.progress(overall_progress / 100)
    st.write(f"Overall Progress: {overall_progress:.2f}%")
    
    # Generate and download project report
    if st.button("Generate Report"):
        report_file = generate_report(project, deadline, progress_data, overall_progress)
        st.success(f"Report saved as {report_file}")
        st.markdown(create_download_link(report_file), unsafe_allow_html=True)
    
    # # Send the report via email
    # st.subheader("Send Report via Email")
    # recipient_emails = st.text_area("Enter recipient emails (comma-separated)")
    # if st.button("Send Report"):
    #     if report_file:
    #         for email in recipient_emails.split(","):
    #             email_status = send_email_with_report(
    #                 email.strip(),
    #                 f"Project Report: {project}",
    #                 f"Dear Team,\n\nPlease find attached the latest project report for {project}.\n\nBest regards,\nProject Management Team",
    #                 report_file,
    #             )
    #             st.write(email_status)
    #     else:
    #         st.error("Please generate the report first!")


# Call the function to display in menu
# project_progress()



# Set milestones for each project
def set_milestones():
    st.header("Set Milestones")
    df = load_projects()
    if df.empty:
        st.info("No projects available.")
        return

    project = st.selectbox("Select Project for Milestone", df["Project Name"])
    milestone = st.text_input("Milestone Name")
    milestone_date = st.date_input("Milestone Date")

    if st.button("Add Milestone"):
        if project and milestone:
            st.success(f"Milestone '{milestone}' added to project '{project}' for {milestone_date}.")
        else:
            st.error("Please complete the milestone details.")

# Time tracking for projects
def time_tracking():
    st.header("Time Tracking")
    project = st.selectbox("Select Project", load_projects()["Project Name"])
    hours_logged = st.number_input("Hours Worked", min_value=0)
    
    if st.button("Log Time"):
        if project:
            st.success(f"{hours_logged} hours logged for project '{project}'.")

# Document uploads
def document_upload():
    st.header("Document Upload")
    project = st.selectbox("Select Project for Document Upload", load_projects()["Project Name"])
    uploaded_file = st.file_uploader("Upload Document", type=["pdf", "docx", "txt"])

    if uploaded_file is not None:
        st.success(f"Document uploaded to project '{project}'.")

# Project tags and categorization
def project_tags():
    st.header("Project Tags & Categories")
    df = load_projects()
    if df.empty:
        st.info("No projects available.")
        return

    project = st.selectbox("Select Project for Tagging", df["Project Name"])
    tags = st.text_input("Tags (comma-separated)")

    if st.button("Add Tags"):
        if tags:
            df.loc[df["Project Name"] == project, "Tags"] = tags
            save_projects(df)
            st.success(f"Tags updated for project '{project}'.")

# Analytics and reports
def analytics():
    st.header("Project Analytics and Reports")
    df = load_projects()
    if df.empty:
        st.info("No data available.")
        return

    st.subheader("Projects by Priority")
    priority_counts = df["Priority"].value_counts()
    st.bar_chart(priority_counts)

    st.subheader("Project Progress Overview")
    st.line_chart(df["Progress"])

# Calendar view
def calendar_view():
    st.header("Project Calendar View")
    df = load_projects()
    if df.empty:
        st.info("No projects to display on the calendar.")
        return

    st.subheader("Upcoming Deadlines")
    for i, row in df.iterrows():
        st.text(f"{row['Project Name']} - End Date: {row['End Date']}")
# Automated notifications for approaching deadlines or high-priority tasks
def automated_notifications():
    st.header("Automated Notifications")
    df = load_projects()
    if df.empty:
        st.info("No projects available.")
        return

    today = date.today()
    notification_list = []

    # Identify upcoming deadlines and high-priority projects
    for _, row in df.iterrows():
        days_left = (pd.to_datetime(row['End Date']) - pd.to_datetime(today)).days
        if days_left <= 7 or row['Priority'] == 'High':
            notification_list.append(f"{row['Project Name']} - {days_left} days remaining")

    if notification_list:
        st.subheader("Notifications")
        for notification in notification_list:
            st.warning(notification)
    else:
        st.info("No upcoming deadlines or high-priority tasks.")


# Project search and advanced filtering
def search_and_filter_projects():
    st.header("Search and Filter Projects")
    df = load_projects()
    if df.empty:
        st.info("No projects available.")
        return

    # Search bar
    search_query = st.text_input("Search by Project Name or Description")

    # Advanced filters
    filter_priority = st.selectbox("Filter by Priority", ["All", "Low", "Medium", "High"])
    filter_status = st.selectbox("Filter by Status", ["All", "Completed", "Uncompleted"])

    # Apply filters and search
    if search_query:
        df = df[df["Project Name"].str.contains(search_query, case=False) | df["Description"].str.contains(search_query, case=False)]
    if filter_priority != "All":
        df = df[df["Priority"] == filter_priority]
    if filter_status != "All":
        df = df[df["Status"] == filter_status]

    if df.empty:
        st.info("No projects matched the search criteria.")
    else:
        st.write(df)

# Backup and restore project data
def backup_and_restore():
    st.header("Backup and Restore")

    # Backup functionality
    if st.button("Backup Data"):
        df = load_projects()
        if not df.empty:
            df.to_csv('project_backup.csv', index=False)
            st.success("Project data backed up successfully.")
        else:
            st.info("No data to back up.")

    # Restore functionality
    uploaded_file = st.file_uploader("Upload Backup File", type=["csv"])
    if uploaded_file is not None:
        restored_df = pd.read_csv(uploaded_file)
        save_projects(restored_df)
        st.success("Project data restored successfully.")

# Customizable status options
def customize_status():
    st.header("Customizable Project Status Options")
    df = load_projects()
    if df.empty:
        st.info("No projects available.")
        return

    project = st.selectbox("Select Project to Change Status", df["Project Name"])
    new_status = st.text_input("Enter New Status (e.g., In Progress, Pending, On Hold)")

    if st.button("Update Status"):
        if project and new_status:
            df.loc[df["Project Name"] == project, "Status"] = new_status
            save_projects(df)
            st.success(f"Status updated for project '{project}' to '{new_status}'.")

# Project templates
def project_templates():
    st.header("Project Templates")
    templates = {
        "Software Development": {
            "Description": "Developing a software application.",
            "Start Date": date.today(),
            "End Date": date.today(),
            "Priority": "High",
            "Status": "Uncompleted"
        },
        "Marketing Campaign": {
            "Description": "Executing a marketing strategy.",
            "Start Date": date.today(),
            "End Date": date.today(),
            "Priority": "Medium",
            "Status": "Uncompleted"
        }
    }

    selected_template = st.selectbox("Choose a Template", list(templates.keys()))
    template_data = templates[selected_template]

    if st.button("Create Project from Template"):
        df = load_projects()
        new_project = pd.DataFrame(template_data, index=[0])
        df = pd.concat([df, new_project], ignore_index=True)
        save_projects(df)
        st.success(f"Project created from template: {selected_template}")

#Budget Tracking
def budget_tracking():
    st.header("Budget Tracking")
    df = load_projects()  # Assume this fetches your existing projects data.

    if df.empty:
        st.info("No projects available.")
        return

    project = st.selectbox("Select Project for Budget Tracking", df["Project Name"])
    budget = st.number_input("Set Budget (in PKR)", min_value=0)
    expenses = st.number_input("Enter Expenses (in PKR)", min_value=0)

    if st.button("Update Budget"):
        # Update the budget and expenses for the selected project.
        df.loc[df["Project Name"] == project, "Budget"] = budget
        df.loc[df["Project Name"] == project, "Expenses"] = expenses
        save_projects(df)
        st.success(f"Budget and expenses updated for project '{project}'.")

    # Display budget status with Profit or Loss Calculation
    st.subheader("Project Budgets")
    for _, row in df.iterrows():
        budget = row.get("Budget", 0)
        expenses = row.get("Expenses", 0)
        remaining_budget = budget - expenses
        # Calculate profit or loss
        if remaining_budget > 0:
            status = "Profit"
            amount = remaining_budget
        else:
            status = "Loss"
            amount = -remaining_budget  # Convert to positive value for loss

        st.write(
            f"{row['Project Name']} - Budget: {budget}, Expenses: {expenses}, "
            f"Remaining: {remaining_budget} ({status}: {amount})"
        )

    # Download options
    st.subheader("Download Budget Reports")

    # Current project budget
    current_project_data = df[df["Project Name"] == project]
    if not current_project_data.empty:
        current_csv = current_project_data.to_csv(index=False)
        st.download_button(
            label=f"Download Current Budget ({project})",
            data=current_csv,
            file_name=f"{project}_budget.csv",
            mime="text/csv",
        )

    # # Overall budgets of the selected project
    # overall_project_data = df[df["Project Name"] == project][["Project Name", "Budget", "Expenses"]]
    # if not overall_project_data.empty:
    #     overall_csv = overall_project_data.to_csv(index=False)
    #     st.download_button(
    #         label=f"Download Overall Budget for {project}",
    #         data=overall_csv,
    #         file_name=f"{project}_overall_budget.csv",
    #         mime="text/csv",
    #     )

    # # All projects' budgets
    # all_csv = df.to_csv(index=False)
    # st.download_button(
    #     label="Download All Projects Budget",
    #     data=all_csv,
    #     file_name="all_projects_budget.csv",
    #     mime="text/csv",
    # )




# Comments and Notes
def comments_and_notes():
    st.header("Comments and Notes")
    df = load_projects()
    if df.empty:
        st.info("No projects available.")
        return

    project = st.selectbox("Select Project to Comment On", df["Project Name"])
    comment = st.text_area("Add Comment/Note")

    if st.button("Add Comment"):
        df.loc[df["Project Name"] == project, "Comments"] = df.loc[df["Project Name"] == project, "Comments"].fillna("") + comment + "\n"
        save_projects(df)
        st.success(f"Comment added to project '{project}'.")

    # Display comments for each project
    st.subheader("Project Comments")
    for _, row in df.iterrows():
        comments = row.get("Comments", "No comments yet.")
        st.write(f"{row['Project Name']} - Comments:\n{comments}")

# Prioritization Metrics
def prioritization_metrics():
    st.header("Prioritization Metrics")
    df = load_projects()
    if df.empty:
        st.info("No projects available.")
        return

    project = st.selectbox("Select Project for Prioritization", df["Project Name"])
    impact = st.slider("Impact", min_value=1, max_value=10)
    cost = st.slider("Cost", min_value=1, max_value=10)
    urgency = st.slider("Urgency", min_value=1, max_value=10)
    overall_priority = (impact * 0.5) + (urgency * 0.3) + (10 - cost) * 0.2

    if st.button("Set Prioritization Metrics"):
        df.loc[df["Project Name"] == project, "Impact"] = impact
        df.loc[df["Project Name"] == project, "Cost"] = cost
        df.loc[df["Project Name"] == project, "Urgency"] = urgency
        df.loc[df["Project Name"] == project, "Priority Score"] = overall_priority
        save_projects(df)
        st.success(f"Prioritization metrics set for project '{project}'.")

    # Display prioritization summary
    st.subheader("Project Prioritization Scores")
    for _, row in df.iterrows():
        priority_score = row.get("Priority Score", "Not set")
        st.write(f"{row['Project Name']} - Priority Score: {priority_score}")

# Calendar View (visualization of deadlines using a calendar layout)
def calendar_view():
    st.header("Calendar View")
    import plotly.express as px
    df = load_projects()
    if df.empty:
        st.info("No projects available.")
        return

    # Convert date columns to datetime
    df["Start Date"] = pd.to_datetime(df["Start Date"])
    df["End Date"] = pd.to_datetime(df["End Date"])

    # Create a Gantt chart to visualize project timelines
    fig = px.timeline(df, x_start="Start Date", x_end="End Date", y="Project Name", color="Priority",
                      title="Project Calendar View", labels={"Priority": "Priority Level"})
    fig.update_yaxes(categoryorder="total ascending")
    st.plotly_chart(fig)
# Automated Notifications for approaching deadlines and task assignments
def automated_notifications():
    st.header("Automated Notifications")
    df = load_projects()
    if df.empty:
        st.info("No projects available.")
        return

    today = date.today()
    notifications = []

    # Check for projects with approaching deadlines
    for _, row in df.iterrows():
        end_date = pd.to_datetime(row["End Date"])
        days_to_deadline = (end_date - today).days
        if days_to_deadline == 7:
            notifications.append(f"Project '{row['Project Name']}' is due in 7 days!")
        elif days_to_deadline == 1:
            notifications.append(f"Project '{row['Project Name']}' is due tomorrow!")
        elif days_to_deadline == 0:
            notifications.append(f"Project '{row['Project Name']}' is due today!")

    # Display notifications
    if notifications:
        st.subheader("Notifications")
        for notification in notifications:
            st.warning(notification)
    else:
        st.info("No notifications for upcoming deadlines.")

# Project Search and Filter by various criteria
def search_and_filter_projects():
    st.header("Search and Filter Projects")
    df = load_projects()
    if df.empty:
        st.info("No projects available.")
        return

    search_term = st.text_input("Search by Project Name or Description")
    priority_filter = st.selectbox("Filter by Priority", ["All", "Low", "Medium", "High"])
    status_filter = st.selectbox("Filter by Status", ["All", "Completed", "Uncompleted"])

    if search_term:
        df = df[df["Project Name"].str.contains(search_term, case=False) | df["Description"].str.contains(search_term, case=False)]
    
    if priority_filter != "All":
        df = df[df["Priority"] == priority_filter]
    
    if status_filter != "All":
        df = df[df["Status"] == status_filter]

    if df.empty:
        st.info("No projects match the search criteria.")
    else:
        st.subheader("Search Results")
        for _, row in df.iterrows():
            st.write(f"**{row['Project Name']}** - Priority: {row['Priority']}, Status: {row['Status']}")

# Backup and Restore project data
def backup_and_restore():
    st.header("Backup and Restore")
    df = load_projects()

    if st.button("Backup Project Data"):
        backup_file = PROJECT_DATA_FILE.replace(".csv", "_backup.csv")
        df.to_csv(backup_file, index=False)
        st.success(f"Backup created as {backup_file}")

    uploaded_file = st.file_uploader("Restore from Backup", type=["csv"])
    if uploaded_file:
        df_backup = pd.read_csv(uploaded_file)
        save_projects(df_backup)
        st.success("Project data restored from backup.")

# Customizable Status Options
def customize_status():
    st.header("Customize Project Status Options")
    df = load_projects()
    if df.empty:
        st.info("No projects available.")
        return

    project = st.selectbox("Select Project to Update Status", df["Project Name"])
    status_options = ["In Progress", "Pending", "On Hold", "Completed", "Uncompleted"]
    custom_status = st.selectbox("Set Custom Status", status_options)

    if st.button("Update Status"):
        df.loc[df["Project Name"] == project, "Status"] = custom_status
        save_projects(df)
        st.success(f"Status of '{project}' updated to '{custom_status}'")

# Project Templates for frequently used project types
def project_templates():
    st.header("Project Templates")
    template_name = st.selectbox("Select Template", ["Software Development", "Marketing Campaign", "Research Project", "Event Planning", "Custom"])
    
    if template_name != "Custom":
        template_description = {
            "Software Development": "Developing a software application with multiple phases, from requirements gathering to deployment.",
            "Marketing Campaign": "Organizing and launching a marketing campaign, including content creation and social media promotion.",
            "Research Project": "Conducting research, including literature review, data collection, analysis, and reporting.",
            "Event Planning": "Planning and organizing an event, including venue booking, logistics, and promotions."
        }.get(template_name, "A general project template.")
    else:
        template_description = st.text_area("Custom Project Description")

    if st.button("Create Project from Template"):
        new_project = pd.DataFrame({
            "Project Name": [template_name + " Project"],
            "Description": [template_description],
            "Start Date": [date.today()],
            "End Date": [date.today()],
            "Priority": ["Medium"],
            "Status": ["Uncompleted"]
        })
        
        df = load_projects()
        df = pd.concat([df, new_project], ignore_index=True)
        save_projects(df)
        st.success(f"Project '{template_name} Project' created from template.")





# Comments and Notes for team collaboration
def comments_and_notes():
    st.header("Project Comments & Notes")
    df = load_projects()
    if df.empty:
        st.info("No projects available.")
        return

    project = st.selectbox("Select Project", df["Project Name"])
    note = st.text_area("Add a Comment or Note")

    if st.button("Add Note"):
        df.loc[df["Project Name"] == project, "Notes"] = df.loc[df["Project Name"] == project].get("Notes", "") + "\n" + note
        save_projects(df)
        st.success(f"Note added to '{project}'")

    st.subheader("Project Notes")
    for i, row in df.iterrows():
        if "Notes" in row and row["Notes"]:
            st.write(f"**{row['Project Name']}** Notes:\n{row['Notes']}")

# Calendar View for visual project timelines and deadlines
def calendar_view():
    st.header("Project Calendar View")
    df = load_projects()
    if df.empty:
        st.info("No projects available.")
        return

    st.write("**Project Deadlines and Timelines**")
    for i, row in df.iterrows():
        st.write(f"{row['Project Name']} - **Start:** {row['Start Date']} **End:** {row['End Date']}")


# Prioritization Metrics for automated prioritization
def prioritization_metrics():
    st.header("Prioritization Metrics")
    df = load_projects()
    if df.empty:
        st.info("No projects available.")
        return

    priority_metric = st.selectbox("Set Prioritization Criteria", ["Impact", "Cost", "Deadline"])
    st.write(f"Projects prioritized by {priority_metric}")

    # Example: prioritize by impact or cost (placeholder logic)
    if priority_metric == "Impact":
        df = df.sort_values("Priority", ascending=False)  # Assuming high = more impact
    elif priority_metric == "Cost":
        df = df.sort_values("Budget", ascending=True)
    elif priority_metric == "Deadline":
        df = df.sort_values("End Date")

    for i, row in df.iterrows():
        st.write(f"{row['Project Name']} - Priority: {row['Priority']} - Deadline: {row['End Date']}")
# Customizable Status Options: Allows users to create custom project statuses
def customizable_status_options():
    st.header("Customizable Status Options")
    df = load_projects()
    
    if df.empty:
        st.info("No projects available.")
        return
    
    # Allow users to select a project and choose a custom status
    project = st.selectbox("Select Project", df["Project Name"].unique())
    custom_status = st.text_input("Enter New Status (e.g., In Progress, On Hold)")
    
    if st.button("Update Status"):
        if custom_status:
            df.loc[df["Project Name"] == project, "Status"] = custom_status
            save_projects(df)
            st.success(f"Status for '{project}' updated to '{custom_status}'.")
        else:
            st.warning("Please enter a valid status.")

# Comments and Notes: Add a comment section to each project
def comments_and_notes():
    st.header("Comments and Notes")
    df = load_projects()
    
    if df.empty:
        st.info("No projects available.")
        return
    
    project = st.selectbox("Select Project for Comments", df["Project Name"].unique())
    comment = st.text_area("Add Comment or Note")
    
    if st.button("Submit Comment"):
        if comment:
            comments_file = f"{project}_comments.txt"
            with open(comments_file, "a") as file:
                file.write(f"{comment}\n")
            st.success("Comment added successfully!")
        else:
            st.warning("Please enter a comment.")
    
    # Display existing comments
    if os.path.exists(f"{project}_comments.txt"):
        st.subheader("Existing Comments")
        with open(f"{project}_comments.txt", "r") as file:
            comments = file.readlines()
            for c in comments:
                st.text(c.strip())

# Automated Notifications for project deadlines and updates
def automated_notifications():
    st.header("Automated Notifications")
    df = load_projects()
    
    if df.empty:
        st.info("No projects available.")
        return

    today = date.today()
    upcoming_deadlines = df[(pd.to_datetime(df['End Date']) >= today) & (pd.to_datetime(df['End Date']) <= today + pd.Timedelta(days=7))]
    
    if not upcoming_deadlines.empty:
        st.warning("Upcoming Deadlines:")
        for i, row in upcoming_deadlines.iterrows():
            st.text(f"{row['Project Name']} - Due on {row['End Date']}")
    
    st.success("Notifications are enabled for approaching deadlines and updates.")

# Team Collaboration: Add team members and assign roles to projects
def team_collaboration():
    st.header("Team Collaboration")
    df = load_projects()
    
    if df.empty:
        st.info("No projects available.")
        return
    
    project = st.selectbox("Select Project for Team Collaboration", df["Project Name"].unique())
    member_name = st.text_input("Team Member Name")
    role = st.selectbox("Role", ["Team Member", "Project Manager", "Viewer"])
    
    if st.button("Add Team Member"):
        if member_name:
            team_file = f"{project}_team.csv"
            if os.path.exists(team_file):
                team_df = pd.read_csv(team_file)
            else:
                team_df = pd.DataFrame(columns=["Member Name", "Role"])
                
            new_member = pd.DataFrame({"Member Name": [member_name], "Role": [role]})
            team_df = pd.concat([team_df, new_member], ignore_index=True)
            team_df.to_csv(team_file, index=False)
            st.success(f"{member_name} added to the project as {role}!")
        else:
            st.warning("Please enter a team member's name.")

    # Display current team members
    if os.path.exists(f"{project}_team.csv"):
        st.subheader("Current Team Members")
        team_df = pd.read_csv(f"{project}_team.csv")
        st.dataframe(team_df)
# Gantt Chart Visualization
def gantt_chart_visualization():
    st.header("Gantt Chart Visualization")
    df = load_projects()
    
    if df.empty:
        st.info("No projects available to display in Gantt chart.")
        return
    
    # Prepare data for Gantt chart
    df['Start Date'] = pd.to_datetime(df['Start Date'])
    df['End Date'] = pd.to_datetime(df['End Date'])
    df['Duration'] = (df['End Date'] - df['Start Date']).dt.days

    chart = alt.Chart(df).mark_bar().encode(
        x='Start Date:T',
        x2='End Date:T',
        y=alt.Y('Project Name', sort=alt.SortField("Start Date", order="ascending")),
        color='Priority'
    ).properties(width=600, height=400)

    st.altair_chart(chart)

# Task Dependencies and Sequencing
def manage_task_dependencies():
    st.header("Manage Task Dependencies")
    project_name = st.selectbox("Select Project", load_projects()["Project Name"].unique())
    task_name = st.text_input("Task Name")
    dependency_task = st.selectbox("Select Prerequisite Task", load_projects()["Project Name"].unique())
    
    if st.button("Set Dependency"):
        if task_name and dependency_task:
            dependencies_file = f"{project_name}_dependencies.csv"
            if os.path.exists(dependencies_file):
                dependencies_df = pd.read_csv(dependencies_file)
            else:
                dependencies_df = pd.DataFrame(columns=["Task", "Dependency"])
            
            new_dependency = pd.DataFrame({"Task": [task_name], "Dependency": [dependency_task]})
            dependencies_df = pd.concat([dependencies_df, new_dependency], ignore_index=True)
            dependencies_df.to_csv(dependencies_file, index=False)
            st.success("Dependency added successfully!")
        else:
            st.warning("Please enter both the task name and prerequisite task.")

    if os.path.exists(f"{project_name}_dependencies.csv"):
        st.subheader("Existing Dependencies")
        dependencies_df = pd.read_csv(f"{project_name}_dependencies.csv")
        st.dataframe(dependencies_df)

# Resource Management
def resource_management():
    st.header("Resource Management")
    df = load_projects()
    if df.empty:
        st.info("No projects available.")
        return
    
    project = st.selectbox("Select Project for Resource Allocation", df["Project Name"].unique())
    resource_name = st.text_input("Resource Name")
    allocation_percentage = st.slider("Resource Allocation Percentage", 0, 100, 50)
    
    if st.button("Allocate Resource"):
        if resource_name:
            resource_file = f"{project}_resources.csv"
            if os.path.exists(resource_file):
                resource_df = pd.read_csv(resource_file)
            else:
                resource_df = pd.DataFrame(columns=["Resource Name", "Allocation (%)"])
            
            new_resource = pd.DataFrame({"Resource Name": [resource_name], "Allocation (%)": [allocation_percentage]})
            resource_df = pd.concat([resource_df, new_resource], ignore_index=True)
            resource_df.to_csv(resource_file, index=False)
            st.success(f"{resource_name} allocated to '{project}'!")
        else:
            st.warning("Please enter the resource name.")

    if os.path.exists(f"{project}_resources.csv"):
        st.subheader("Allocated Resources")
        resource_df = pd.read_csv(f"{project}_resources.csv")
        st.dataframe(resource_df)

# Risk Management and Mitigation Tracking
def risk_management():
    st.header("Risk Management and Mitigation Tracking")
    df = load_projects()
    
    if df.empty:
        st.info("No projects available.")
        return

    project = st.selectbox("Select Project for Risk Management", df["Project Name"].unique())
    risk_description = st.text_input("Risk Description")
    risk_level = st.selectbox("Risk Level", ["Low", "Medium", "High"])
    mitigation_strategy = st.text_area("Mitigation Strategy")
    
    if st.button("Log Risk"):
        if risk_description and mitigation_strategy:
            risk_file = f"{project}_risks.csv"
            if os.path.exists(risk_file):
                risk_df = pd.read_csv(risk_file)
            else:
                risk_df = pd.DataFrame(columns=["Risk Description", "Risk Level", "Mitigation Strategy"])
            
            new_risk = pd.DataFrame({
                "Risk Description": [risk_description],
                "Risk Level": [risk_level],
                "Mitigation Strategy": [mitigation_strategy]
            })
            risk_df = pd.concat([risk_df, new_risk], ignore_index=True)
            risk_df.to_csv(risk_file, index=False)
            st.success(f"Risk added to '{project}'!")
        else:
            st.warning("Please complete all fields to log a risk.")

    if os.path.exists(f"{project}_risks.csv"):
        st.subheader("Logged Risks")
        risk_df = pd.read_csv(f"{project}_risks.csv")
        st.dataframe(risk_df)

# Integration with Third-Party Tools (Simulated example)
def integration_with_third_party_tools():
    st.header("Integration with Third-Party Tools")
    st.info("Here, you can integrate with various third-party tools like Slack, Trello, or Google Workspace.")

    tool = st.selectbox("Select Tool to Integrate", ["Slack", "GitHub", "Trello", "Google Workspace"])
    if st.button("Integrate"):
        st.success(f"{tool} has been integrated successfully! (Simulated)")



  
# #project tags and categories , "Third-Party Integrations",  "Automated Notifications", "Backup & Restore", 


# for notifications:

import streamlit as st

# Add your other function definitions here, like add_project(), edit_project(), etc.

# Sidebar menu using radio buttons
menu = st.sidebar.radio(
    "Menu", [
        "Home", "Add Project", "Edit Project", "Delete Project", 
        "Task Management", "High Priority & Deadlines", 
         "Project Progress Tracking", 
          
        "Analytics & Reports", "Calendar View", 
        "Search and Filter", "Budget Tracking", 
        "Comments and Notes", "Prioritization Metrics", 
        "Gantt Chart Visualization", 
        "Resource Management", "Risk Management"
    ]
)

# Linking each radio button option to the corresponding function
if menu == "Home":
    home()
elif menu == "Add Project":
    add_project()
elif menu == "Edit Project":
    edit_project()
elif menu == "Delete Project":
    delete_project()
elif menu == "Task Management":
    manage_tasks()  # Calling your manage_tasks function here
elif menu == "High Priority & Deadlines":
    view_high_priority()
# elif menu == "Team Collaboration":
#     manage_team()
elif menu == "Project Progress Tracking":
    project_progress()
# elif menu == "Set Milestones":
#     set_milestones()
# elif menu == "Time Tracking":
#     time_tracking()
# elif menu == "Document Upload":
#     document_upload()
elif menu == "Analytics & Reports":
    analytics()
elif menu == "Calendar View":
    calendar_view()
elif menu == "Search and Filter":
    search_and_filter_projects()
# elif menu == "Customizable Status":
#     customizable_status_options()
# elif menu == "Project Templates":
#     project_templates()
elif menu == "Budget Tracking":
    budget_tracking()
elif menu == "Comments and Notes":
    comments_and_notes()
elif menu == "Prioritization Metrics":
    prioritization_metrics()
elif menu == "Gantt Chart Visualization":
    gantt_chart_visualization()
# elif menu == "Task Dependencies and Sequencing":
#     manage_task_dependencies()
elif menu == "Resource Management":
    resource_management()
elif menu == "Risk Management":
    risk_management()



# "Team Collaboration", "Set Milestones", "Time Tracking", "Document Upload", "Customizable Status", 
        # "Project Templates", "Task Dependencies and Sequencing", 
